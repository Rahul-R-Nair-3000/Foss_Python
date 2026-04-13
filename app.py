# ============================================================
#  AI Resume Analyzer — app.py
#  FOSSEE Internship Screening Project
#  Author  : (your name here)
#  Stack   : Python · Streamlit · pdfplumber · python-docx
# ============================================================

import io
import re
import json
import time
import base64
import textwrap
from collections import Counter
from datetime import datetime

import streamlit as st
import pdfplumber
import docx
import plotly.graph_objects as go
import plotly.express as px
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
)

# ──────────────────────────────────────────────────────────────
#  PAGE CONFIG  (must be first Streamlit call)
# ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="AI Resume Analyzer",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────────────────────
#  GLOBAL CSS  — refined dark-teal editorial theme
# ──────────────────────────────────────────────────────────────

CUSTOM_CSS = """
<style>
/* ── Google Fonts ── */
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500&display=swap');

/* ── Root palette ── */
:root {
    --bg:        #0d1117;
    --surface:   #161b22;
    --surface2:  #1c2330;
    --border:    #30363d;
    --accent:    #2dd4bf;   /* teal */
    --accent2:   #f59e0b;   /* amber */
    --danger:    #f87171;
    --ok:        #4ade80;
    --txt:       #e6edf3;
    --txt-muted: #8b949e;
    --radius:    12px;
}

/* ── Global reset ── */
html, body, [class*="css"] {
    font-family: 'DM Sans', sans-serif;
    background-color: var(--bg) !important;
    color: var(--txt) !important;
}

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border);
}
section[data-testid="stSidebar"] * { color: var(--txt) !important; }

/* ── Hide default Streamlit chrome ── */
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding-top: 2rem !important; }

/* ── Typography ── */
h1, h2, h3, .display-font {
    font-family: 'Syne', sans-serif !important;
}

/* ── Card ── */
.card {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1.4rem 1.6rem;
    margin-bottom: 1rem;
}
.card-accent {
    border-left: 4px solid var(--accent);
}

/* ── Score ring label ── */
.score-big {
    font-family: 'Syne', sans-serif;
    font-size: 3.6rem;
    font-weight: 800;
    color: var(--accent);
    line-height: 1;
}
.score-label {
    font-size: 0.8rem;
    color: var(--txt-muted);
    letter-spacing: 0.08em;
    text-transform: uppercase;
}

/* ── Tag / pill ── */
.tag {
    display: inline-block;
    background: rgba(45,212,191,0.12);
    color: var(--accent);
    border: 1px solid rgba(45,212,191,0.3);
    border-radius: 999px;
    padding: 3px 12px;
    font-size: 0.78rem;
    margin: 3px 2px;
    font-weight: 500;
}
.tag-amber {
    background: rgba(245,158,11,0.12);
    color: var(--accent2);
    border-color: rgba(245,158,11,0.3);
}
.tag-red {
    background: rgba(248,113,113,0.12);
    color: var(--danger);
    border-color: rgba(248,113,113,0.3);
}
.tag-green {
    background: rgba(74,222,128,0.12);
    color: var(--ok);
    border-color: rgba(74,222,128,0.3);
}

/* ── Section heading strip ── */
.section-heading {
    font-family: 'Syne', sans-serif;
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: var(--accent);
    padding-bottom: 6px;
    border-bottom: 1px solid var(--border);
    margin-bottom: 1rem;
}

/* ── Suggestion list ── */
.suggestion-item {
    display: flex;
    gap: 10px;
    align-items: flex-start;
    padding: 0.55rem 0;
    border-bottom: 1px solid var(--border);
    font-size: 0.9rem;
}
.suggestion-item:last-child { border-bottom: none; }
.sug-icon { font-size: 1.1rem; flex-shrink: 0; }

/* ── Progress bar override ── */
div[data-testid="stProgress"] > div > div > div {
    background: var(--accent) !important;
}

/* ── File uploader ── */
section[data-testid="stFileUploader"] {
    background: var(--surface2) !important;
    border: 2px dashed var(--border) !important;
    border-radius: var(--radius) !important;
}

/* ── Metric labels ── */
[data-testid="metric-container"] label {
    color: var(--txt-muted) !important;
    font-size: 0.75rem !important;
    text-transform: uppercase;
    letter-spacing: 0.05em;
}
[data-testid="metric-container"] [data-testid="stMetricValue"] {
    font-family: 'Syne', sans-serif !important;
    font-size: 2rem !important;
    color: var(--txt) !important;
}

/* ── Scrollable text box ── */
.scrollbox {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1rem 1.2rem;
    max-height: 340px;
    overflow-y: auto;
    font-size: 0.82rem;
    line-height: 1.7;
    color: var(--txt-muted);
    white-space: pre-wrap;
}

/* ── Expander ── */
details summary {
    font-family: 'Syne', sans-serif;
    font-weight: 600;
}

/* ── Download button ── */
.dl-btn a {
    display: inline-block;
    background: var(--accent);
    color: #0d1117 !important;
    font-family: 'Syne', sans-serif;
    font-weight: 700;
    font-size: 0.85rem;
    padding: 0.55rem 1.4rem;
    border-radius: 999px;
    text-decoration: none;
    letter-spacing: 0.04em;
    transition: opacity 0.2s;
}
.dl-btn a:hover { opacity: 0.85; }

/* ── Hero banner ── */
.hero {
    background: linear-gradient(135deg, #0f2027, #1a3a3a, #0d1117);
    border: 1px solid var(--border);
    border-radius: 16px;
    padding: 2.5rem 2rem;
    margin-bottom: 2rem;
    position: relative;
    overflow: hidden;
}
.hero::after {
    content: '';
    position: absolute;
    top: -60px; right: -60px;
    width: 220px; height: 220px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(45,212,191,0.18) 0%, transparent 70%);
    pointer-events: none;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: 2.4rem;
    font-weight: 800;
    color: var(--txt);
    margin: 0 0 0.4rem;
}
.hero-sub {
    color: var(--txt-muted);
    font-size: 1rem;
    margin: 0;
}
</style>
"""

st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────
#  SKILL & KEYWORD TAXONOMY
# ──────────────────────────────────────────────────────────────

SKILLS_DB = {
    "Programming Languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "c",
        "rust", "go", "kotlin", "swift", "ruby", "php", "r", "matlab",
        "scala", "perl", "dart", "lua", "haskell",
    ],
    "Web & Frontend": [
        "html", "css", "react", "angular", "vue", "next.js", "svelte",
        "tailwind", "bootstrap", "jquery", "sass", "webpack", "vite",
    ],
    "Backend & APIs": [
        "node.js", "django", "flask", "fastapi", "spring boot", "express",
        "rest api", "graphql", "grpc", "asp.net",
    ],
    "Data & ML": [
        "machine learning", "deep learning", "tensorflow", "pytorch",
        "scikit-learn", "pandas", "numpy", "matplotlib", "seaborn",
        "nlp", "computer vision", "data analysis", "statistics",
        "power bi", "tableau", "excel",
    ],
    "Databases": [
        "sql", "mysql", "postgresql", "mongodb", "redis", "sqlite",
        "elasticsearch", "firebase", "dynamodb", "oracle",
    ],
    "DevOps & Cloud": [
        "docker", "kubernetes", "aws", "azure", "gcp", "ci/cd",
        "jenkins", "github actions", "terraform", "ansible", "linux",
        "bash", "git",
    ],
    "Soft Skills": [
        "leadership", "communication", "teamwork", "problem solving",
        "project management", "agile", "scrum", "time management",
    ],
}

SECTION_KEYWORDS = {
    "education":    ["education", "degree", "university", "college", "school",
                     "bachelor", "master", "phd", "b.tech", "m.tech", "b.e",
                     "b.sc", "diploma", "cgpa", "gpa", "academic"],
    "experience":   ["experience", "internship", "work history", "employment",
                     "job", "position", "role", "responsibilities", "worked",
                     "company", "organization", "tenure"],
    "skills":       ["skills", "technical skills", "technologies", "tools",
                     "proficiencies", "expertise", "competencies", "stack"],
    "projects":     ["project", "projects", "built", "developed", "implemented",
                     "created", "github", "portfolio", "capstone"],
    "achievements": ["achievement", "award", "honor", "certification",
                     "certificate", "hackathon", "competition", "rank",
                     "distinction", "recognition"],
    "summary":      ["summary", "objective", "about me", "profile",
                     "introduction", "overview", "career goal"],
    "contact":      ["email", "phone", "linkedin", "github", "twitter",
                     "address", "contact", "mobile", "website"],
}

POWER_VERBS = [
    "achieved", "built", "created", "designed", "developed", "engineered",
    "established", "executed", "implemented", "improved", "increased",
    "launched", "led", "managed", "optimized", "orchestrated", "reduced",
    "spearheaded", "streamlined", "transformed", "delivered", "automated",
]

# ──────────────────────────────────────────────────────────────
#  FILE PARSING
# ──────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF using pdfplumber."""
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as exc:
        raise ValueError(f"Could not read PDF: {exc}") from exc
    return "\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # also pull table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        raise ValueError(f"Could not read DOCX: {exc}") from exc

# ──────────────────────────────────────────────────────────────
#  TEXT ANALYSIS
# ──────────────────────────────────────────────────────────────

def extract_contact_info(text: str) -> dict:
    """Pull email, phone and URLs from raw text."""
    email = re.findall(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text)
    phone = re.findall(
        r"(\+?\d[\d\s\-().]{7,}\d)", text
    )
    urls  = re.findall(
        r"(https?://[^\s]+|linkedin\.com/in/[^\s]+|github\.com/[^\s]+)", text
    )
    return {
        "email":  list(set(email))[:2],
        "phone":  [p.strip() for p in list(set(phone))[:2]],
        "urls":   list(set(urls))[:5],
    }


def detect_sections(text: str) -> dict:
    """Return which major resume sections are present."""
    lower = text.lower()
    return {
        section: any(kw in lower for kw in keywords)
        for section, keywords in SECTION_KEYWORDS.items()
    }


def extract_skills(text: str) -> dict:
    """Match skills from the taxonomy against resume text."""
    lower = text.lower()
    found: dict[str, list[str]] = {}
    for category, skills in SKILLS_DB.items():
        matched = [s for s in skills if re.search(rf"\b{re.escape(s)}\b", lower)]
        if matched:
            found[category] = matched
    return found


def count_power_verbs(text: str) -> list[str]:
    """Return action verbs present in the resume."""
    lower = text.lower()
    return [v for v in POWER_VERBS if re.search(rf"\b{v}\b", lower)]


def word_count(text: str) -> int:
    return len(text.split())


def estimate_pages(text: str) -> float:
    """Very rough: ~500 words per page."""
    return round(word_count(text) / 500, 1)

# ──────────────────────────────────────────────────────────────
#  SCORING ENGINE
# ──────────────────────────────────────────────────────────────

def compute_score(text: str, sections: dict, skills: dict) -> dict:
    """
    Returns a breakdown dict and total score (0–100).
    Categories:
      • Sections completeness  — 30 pts
      • Skills breadth         — 25 pts
      • Contact info           — 15 pts
      • Power verbs / writing  — 15 pts
      • Length / detail        — 15 pts
    """
    # 1. Sections (6 key ones × 5 pts each = 30)
    key_sections = ["education", "experience", "skills", "projects",
                    "achievements", "summary"]
    section_score = sum(5 for s in key_sections if sections.get(s, False))

    # 2. Skills (max 25)
    total_skills = sum(len(v) for v in skills.values())
    skill_score = min(25, total_skills * 2)

    # 3. Contact (max 15)
    contact = extract_contact_info(text)
    contact_score = 0
    if contact["email"]:  contact_score += 6
    if contact["phone"]:  contact_score += 4
    if contact["urls"]:   contact_score += 5

    # 4. Power verbs (max 15)
    verbs = count_power_verbs(text)
    verb_score = min(15, len(verbs) * 2)

    # 5. Length (max 15)
    wc = word_count(text)
    if   wc >= 500: length_score = 15
    elif wc >= 300: length_score = 10
    elif wc >= 150: length_score = 5
    else:           length_score = 0

    total = section_score + skill_score + contact_score + verb_score + length_score

    return {
        "total": total,
        "breakdown": {
            "Sections":     (section_score, 30),
            "Skills":       (skill_score,   25),
            "Contact Info": (contact_score, 15),
            "Action Verbs": (verb_score,    15),
            "Detail/Length":(length_score,  15),
        },
        "contact": contact,
        "power_verbs": verbs,
    }


def grade_label(score: int) -> tuple[str, str]:
    """Return (grade letter, colour class) for a score 0-100."""
    if score >= 85: return "A+", "tag-green"
    if score >= 70: return "A",  "tag-green"
    if score >= 55: return "B",  "tag-amber"
    if score >= 40: return "C",  "tag-amber"
    return "D", "tag-red"

# ──────────────────────────────────────────────────────────────
#  SUGGESTIONS ENGINE
# ──────────────────────────────────────────────────────────────

def generate_suggestions(text: str, sections: dict, skills: dict,
                          score_data: dict) -> list[dict]:
    """Return a list of {priority, icon, text} suggestion dicts."""
    suggestions = []
    bd = score_data["breakdown"]

    # Missing sections
    missing_sections = [s for s in ["summary", "experience", "education",
                                     "skills", "projects", "achievements"]
                        if not sections.get(s)]
    if missing_sections:
        suggestions.append({
            "priority": "high",
            "icon": "🔴",
            "text": f"Add missing section(s): **{', '.join(s.title() for s in missing_sections)}**. "
                    "Recruiters scan for these headings in seconds.",
        })

    # Too short
    wc = word_count(text)
    if wc < 300:
        suggestions.append({
            "priority": "high",
            "icon": "🔴",
            "text": f"Your resume has only **{wc} words** — expand it to at least 400–600 words with "
                    "more detail about your roles, projects, and responsibilities.",
        })
    elif wc > 1200:
        suggestions.append({
            "priority": "medium",
            "icon": "🟡",
            "text": f"Your resume is **{wc} words** — consider trimming to 1–2 pages. "
                    "Conciseness signals clarity of thought.",
        })

    # Skill diversity
    total_skills = sum(len(v) for v in skills.values())
    if total_skills < 5:
        suggestions.append({
            "priority": "high",
            "icon": "🔴",
            "text": "Only **few skills detected**. List your technical stack, tools, and frameworks "
                    "explicitly so ATS (Applicant Tracking Systems) can find them.",
        })
    missing_cats = [c for c in ["Programming Languages", "Databases", "DevOps & Cloud"]
                    if c not in skills]
    if missing_cats:
        suggestions.append({
            "priority": "medium",
            "icon": "🟡",
            "text": f"Consider adding skills from: **{', '.join(missing_cats)}** "
                    "if they apply to your experience.",
        })

    # Power verbs
    if bd["Action Verbs"][0] < 6:
        suggestions.append({
            "priority": "medium",
            "icon": "🟡",
            "text": "Use more **action/power verbs** (e.g., *engineered, optimised, spearheaded, delivered*) "
                    "to make your bullet points punchy and impactful.",
        })

    # Contact
    contact = score_data["contact"]
    if not contact["email"]:
        suggestions.append({"priority": "high", "icon": "🔴",
                             "text": "**No email address detected.** Add professional contact details."})
    if not contact["urls"]:
        suggestions.append({"priority": "medium", "icon": "🟡",
                             "text": "Add your **LinkedIn** and/or **GitHub** URL — most tech recruiters "
                                     "visit profiles before interviews."})

    # Quantification
    numbers = re.findall(r"\b\d+[\d,%+x]*\b", text)
    if len(numbers) < 4:
        suggestions.append({
            "priority": "medium",
            "icon": "🟡",
            "text": "Quantify your achievements: replace vague statements with numbers "
                    "(e.g., *'Reduced load time by 40%'*, *'Led team of 5'*).",
        })

    # Positive notes
    if bd["Sections"][0] >= 25:
        suggestions.append({"priority": "good", "icon": "✅",
                             "text": "Great section coverage — all key resume headings are present!"})
    if total_skills >= 10:
        suggestions.append({"priority": "good", "icon": "✅",
                             "text": f"Solid skills profile with **{total_skills} skills** detected across categories."})
    if bd["Action Verbs"][0] >= 10:
        suggestions.append({"priority": "good", "icon": "✅",
                             "text": "Strong use of action verbs throughout your resume."})

    return suggestions

# ──────────────────────────────────────────────────────────────
#  CHARTS
# ──────────────────────────────────────────────────────────────

PLOTLY_THEME = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(family="DM Sans", color="#8b949e", size=12),
)


def chart_skill_distribution(skills: dict) -> go.Figure:
    categories = list(skills.keys())
    counts = [len(v) for v in skills.values()]
    fig = go.Figure(go.Bar(
        x=counts, y=categories,
        orientation="h",
        marker=dict(
            color=counts,
            colorscale=[[0, "#1c2330"], [0.5, "#2dd4bf"], [1, "#f59e0b"]],
            showscale=False,
        ),
        text=counts, textposition="outside",
        textfont=dict(color="#e6edf3"),
    ))
    fig.update_layout(
        **PLOTLY_THEME,
        height=max(220, len(categories) * 48),
        margin=dict(l=0, r=40, t=10, b=10),
        xaxis=dict(showgrid=False, zeroline=False, showticklabels=False),
        yaxis=dict(showgrid=False, tickfont=dict(color="#e6edf3", size=12)),
    )
    return fig


def chart_score_gauge(score: int) -> go.Figure:
    grade, _ = grade_label(score)
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        domain={"x": [0, 1], "y": [0, 1]},
        title={"text": f"Resume Score · Grade {grade}",
               "font": {"size": 14, "color": "#8b949e", "family": "DM Sans"}},
        number={"font": {"size": 48, "color": "#2dd4bf", "family": "Syne"},
                "suffix": "/100"},
        gauge={
            "axis":      {"range": [0, 100], "tickcolor": "#30363d",
                          "tickfont": {"color": "#8b949e", "size": 10}},
            "bar":       {"color": "#2dd4bf", "thickness": 0.28},
            "bgcolor":   "#161b22",
            "bordercolor": "#30363d",
            "steps": [
                {"range": [0,  40],  "color": "#1c1020"},
                {"range": [40, 70],  "color": "#1c2218"},
                {"range": [70, 100], "color": "#162520"},
            ],
            "threshold": {"line": {"color": "#f59e0b", "width": 3},
                          "thickness": 0.85, "value": 70},
        },
    ))
    fig.update_layout(
        **PLOTLY_THEME,
        height=260,
        margin=dict(l=30, r=30, t=40, b=10),
    )
    return fig


def chart_score_radar(breakdown: dict) -> go.Figure:
    cats  = list(breakdown.keys())
    vals  = [v[0] / v[1] * 100 for v in breakdown.values()]
    cats += [cats[0]]
    vals += [vals[0]]
    fig = go.Figure(go.Scatterpolar(
        r=vals, theta=cats,
        fill="toself",
        fillcolor="rgba(45,212,191,0.15)",
        line=dict(color="#2dd4bf", width=2),
        marker=dict(size=6, color="#2dd4bf"),
    ))
    fig.update_layout(
        **PLOTLY_THEME,
        polar=dict(
            bgcolor="rgba(0,0,0,0)",
            radialaxis=dict(range=[0, 100], tickfont=dict(color="#8b949e", size=9),
                            gridcolor="#30363d", linecolor="#30363d"),
            angularaxis=dict(tickfont=dict(color="#e6edf3", size=11),
                             gridcolor="#30363d", linecolor="#30363d"),
        ),
        height=300,
        margin=dict(l=30, r=30, t=20, b=20),
    )
    return fig

# ──────────────────────────────────────────────────────────────
#  PDF REPORT GENERATION  (download feature)
# ──────────────────────────────────────────────────────────────

def generate_pdf_report(filename: str, score_data: dict, sections: dict,
                         skills: dict, suggestions: list) -> bytes:
    """Build and return a PDF report as bytes using ReportLab."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             leftMargin=2*cm, rightMargin=2*cm,
                             topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()

    teal    = colors.HexColor("#2dd4bf")
    dark    = colors.HexColor("#0d1117")
    muted   = colors.HexColor("#555555")
    danger  = colors.HexColor("#e05252")
    ok_col  = colors.HexColor("#3a9e5a")
    amber   = colors.HexColor("#c97c10")

    title_style = ParagraphStyle("Title", fontName="Helvetica-Bold",
                                 fontSize=20, textColor=teal, spaceAfter=4)
    sub_style   = ParagraphStyle("Sub", fontName="Helvetica",
                                 fontSize=10, textColor=muted, spaceAfter=12)
    heading_sty = ParagraphStyle("H2", fontName="Helvetica-Bold",
                                 fontSize=12, textColor=dark, spaceBefore=10, spaceAfter=4)
    body_sty    = ParagraphStyle("Body", fontName="Helvetica",
                                 fontSize=9, textColor=colors.HexColor("#333333"),
                                 leading=13)

    story = []

    # Header
    story.append(Paragraph("AI Resume Analyzer — Report", title_style))
    story.append(Paragraph(
        f"File: <b>{filename}</b> &nbsp;·&nbsp; Generated: {datetime.now().strftime('%d %b %Y %H:%M')}",
        sub_style))
    story.append(HRFlowable(width="100%", thickness=1, color=teal))
    story.append(Spacer(1, 10))

    # Score summary table
    grade, _ = grade_label(score_data["total"])
    score_table_data = [
        ["Overall Score", "Grade", "Word Count"],
        [str(score_data["total"]) + " / 100", grade,
         str(word_count(""))],  # placeholder — we don't carry text here
    ]
    t = Table([
        [Paragraph("Overall Score", heading_sty),
         Paragraph("Grade", heading_sty)],
        [Paragraph(f"{score_data['total']} / 100", ParagraphStyle(
            "Big", fontName="Helvetica-Bold", fontSize=22, textColor=teal)),
         Paragraph(grade, ParagraphStyle(
             "Grade", fontName="Helvetica-Bold", fontSize=22, textColor=dark))],
    ], colWidths=[8*cm, 8*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#e8faf8")),
        ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#cccccc")),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white]),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    # Breakdown
    story.append(Paragraph("Score Breakdown", heading_sty))
    bd_rows = [["Category", "Score", "Max"]]
    for cat, (got, mx) in score_data["breakdown"].items():
        bd_rows.append([cat, str(got), str(mx)])
    bt = Table(bd_rows, colWidths=[10*cm, 3*cm, 3*cm])
    bt.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0), colors.HexColor("#f0f0f0")),
        ("FONTNAME",      (0,0), (-1,0), "Helvetica-Bold"),
        ("GRID",          (0,0), (-1,-1), 0.4, colors.HexColor("#dddddd")),
        ("ALIGN",         (1,0), (-1,-1), "CENTER"),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [colors.white, colors.HexColor("#fafafa")]),
    ]))
    story.append(bt)
    story.append(Spacer(1, 12))

    # Sections detected
    story.append(Paragraph("Sections Detected", heading_sty))
    sec_rows = [["Section", "Present"]]
    for sec, present in sections.items():
        sec_rows.append([sec.title(), "✔ Yes" if present else "✘ Missing"])
    st2 = Table(sec_rows, colWidths=[10*cm, 6*cm])
    st2.setStyle(TableStyle([
        ("BACKGROUND",    (0,0), (-1,0), colors.HexColor("#f0f0f0")),
        ("FONTNAME",      (0,0), (-1,0), "Helvetica-Bold"),
        ("GRID",          (0,0), (-1,-1), 0.4, colors.HexColor("#dddddd")),
        ("ALIGN",         (1,0), (-1,-1), "CENTER"),
        ("ROWBACKGROUNDS",(0,1), (-1,-1), [colors.white, colors.HexColor("#fafafa")]),
    ]))
    story.append(st2)
    story.append(Spacer(1, 12))

    # Skills
    story.append(Paragraph("Detected Skills", heading_sty))
    for cat, skill_list in skills.items():
        story.append(Paragraph(
            f"<b>{cat}</b>: {', '.join(skill_list)}", body_sty))
    story.append(Spacer(1, 12))

    # Suggestions
    story.append(Paragraph("Improvement Suggestions", heading_sty))
    for idx, sug in enumerate(suggestions, 1):
        clean = re.sub(r"\*\*(.*?)\*\*", r"\1", sug["text"])  # strip markdown bold
        color = ok_col if sug["priority"] == "good" else (
                danger if sug["priority"] == "high" else amber)
        story.append(Paragraph(
            f"<font color='#{color.hexval()[1:]}'>{sug['icon']}</font> {clean}",
            body_sty))
        story.append(Spacer(1, 4))

    story.append(Spacer(1, 16))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#cccccc")))
    story.append(Paragraph("Generated by AI Resume Analyzer · FOSSEE Internship Project",
                            ParagraphStyle("Footer", fontName="Helvetica",
                                           fontSize=8, textColor=muted, spaceBefore=6)))
    doc.build(story)
    return buf.getvalue()


def get_download_link(data: bytes, filename: str, label: str) -> str:
    b64 = base64.b64encode(data).decode()
    return (f'<div class="dl-btn">'
            f'<a href="data:application/pdf;base64,{b64}" download="{filename}">'
            f'{label}</a></div>')

# ──────────────────────────────────────────────────────────────
#  SIDEBAR
# ──────────────────────────────────────────────────────────────

def render_sidebar():
    with st.sidebar:
        st.markdown("""
        <div style="padding:1.2rem 0 0.5rem">
          <p style="font-family:'Syne',sans-serif;font-size:1.3rem;
                    font-weight:800;color:#2dd4bf;margin:0;">📄 ResumeAI</p>
          <p style="font-size:0.75rem;color:#8b949e;margin:2px 0 0;">
            FOSSEE Internship Screening
          </p>
        </div>
        <hr style="border-color:#30363d;margin:0.6rem 0 1rem">
        """, unsafe_allow_html=True)

        page = st.radio(
            "Navigate",
            ["🏠  Analyzer", "📖  How It Works", "ℹ️  About"],
            label_visibility="collapsed",
        )

        st.markdown("<hr style='border-color:#30363d;margin:1rem 0'>",
                    unsafe_allow_html=True)

        st.markdown("""
        <div style='font-size:0.75rem;color:#8b949e;line-height:1.8'>
          <b style='color:#e6edf3'>Supported formats</b><br>
          PDF &nbsp;·&nbsp; DOCX<br><br>
          <b style='color:#e6edf3'>Scoring weights</b><br>
          Sections &nbsp;30 pts<br>
          Skills &nbsp;&nbsp;&nbsp;&nbsp;&nbsp;25 pts<br>
          Contact &nbsp;&nbsp;15 pts<br>
          Verbs &nbsp;&nbsp;&nbsp;&nbsp;15 pts<br>
          Detail &nbsp;&nbsp;&nbsp;&nbsp;15 pts
        </div>
        """, unsafe_allow_html=True)

    return page

# ──────────────────────────────────────────────────────────────
#  PAGE: ANALYZER
# ──────────────────────────────────────────────────────────────

def page_analyzer():
    # Hero banner
    st.markdown("""
    <div class="hero">
      <p class="hero-title">AI Resume Analyzer</p>
      <p class="hero-sub">
        Upload your resume and get instant insights — skills, sections,
        score &amp; actionable improvement tips.
      </p>
    </div>
    """, unsafe_allow_html=True)

    # ── Upload widget ──
    uploaded = st.file_uploader(
        "Drop your resume here",
        type=["pdf", "docx"],
        help="PDF or DOCX file, max 10 MB",
        label_visibility="collapsed",
    )

    if uploaded is None:
        st.markdown("""
        <div class="card" style="text-align:center;padding:2.5rem 1rem;color:#8b949e;">
          <div style="font-size:3rem;margin-bottom:0.5rem">📂</div>
          <div style="font-family:'Syne',sans-serif;font-size:1.1rem;
                      font-weight:700;color:#e6edf3;">
            Upload a PDF or DOCX resume to begin
          </div>
          <div style="font-size:0.85rem;margin-top:0.4rem">
            Your file is processed locally — nothing is stored.
          </div>
        </div>
        """, unsafe_allow_html=True)
        return

    # ── Process ──
    file_bytes = uploaded.read()

    # Basic validation
    if len(file_bytes) == 0:
        st.error("⚠️ The uploaded file appears to be empty. Please try another file.")
        return
    if len(file_bytes) > 10 * 1024 * 1024:
        st.error("⚠️ File exceeds 10 MB. Please upload a smaller resume.")
        return

    with st.spinner("🔍 Analyzing your resume…"):
        time.sleep(0.4)   # brief pause so spinner is visible
        try:
            ext = uploaded.name.rsplit(".", 1)[-1].lower()
            if ext == "pdf":
                text = extract_text_from_pdf(file_bytes)
            elif ext in ("doc", "docx"):
                text = extract_text_from_docx(file_bytes)
            else:
                st.error("Unsupported file type. Please upload PDF or DOCX.")
                return
        except ValueError as exc:
            st.error(f"❌ Could not parse file: {exc}")
            return

        if not text.strip():
            st.error("⚠️ No readable text found. The file may be image-only or corrupted.")
            return

        sections    = detect_sections(text)
        skills      = extract_skills(text)
        score_data  = compute_score(text, sections, skills)
        suggestions = generate_suggestions(text, sections, skills, score_data)

    # ── Top metrics bar ──
    grade, grade_cls = grade_label(score_data["total"])
    wc = word_count(text)
    pg = estimate_pages(text)
    total_skills = sum(len(v) for v in skills.values())
    sections_found = sum(1 for v in sections.values() if v)

    m1, m2, m3, m4 = st.columns(4)
    with m1:
        st.metric("Resume Score", f"{score_data['total']} / 100")
    with m2:
        st.metric("Grade", grade)
    with m3:
        st.metric("Skills Found", total_skills)
    with m4:
        st.metric("Sections Detected", f"{sections_found} / {len(sections)}")

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Tabs ──
    tab1, tab2, tab3, tab4, tab5 = st.tabs(
        ["📊 Score & Charts", "🛠 Skills", "📋 Sections",
         "💡 Suggestions", "📝 Raw Text"]
    )

    # ── Tab 1: Score & Charts ──
    with tab1:
        c1, c2 = st.columns([1, 1], gap="large")

        with c1:
            st.markdown('<div class="section-heading">Overall Score</div>',
                        unsafe_allow_html=True)
            st.plotly_chart(chart_score_gauge(score_data["total"]),
                            use_container_width=True)

        with c2:
            st.markdown('<div class="section-heading">Dimension Radar</div>',
                        unsafe_allow_html=True)
            st.plotly_chart(chart_score_radar(score_data["breakdown"]),
                            use_container_width=True)

        st.markdown('<div class="section-heading">Score Breakdown</div>',
                    unsafe_allow_html=True)

        for cat, (got, mx) in score_data["breakdown"].items():
            pct = int(got / mx * 100)
            st.markdown(
                f'<div style="display:flex;justify-content:space-between;'
                f'font-size:0.85rem;margin-bottom:2px;">'
                f'<span>{cat}</span>'
                f'<span style="color:#2dd4bf;font-weight:600;">{got}/{mx}</span></div>',
                unsafe_allow_html=True)
            st.progress(pct)
            st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)

        # Contact info
        contact = score_data["contact"]
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-heading">Contact Info Detected</div>',
                    unsafe_allow_html=True)

        def tag(text, cls="tag"):
            return f'<span class="{cls}">{text}</span>'

        contact_html = ""
        for e in contact["email"]:
            contact_html += tag(f"✉ {e}")
        for p in contact["phone"]:
            contact_html += tag(f"📞 {p}", "tag tag-amber")
        for u in contact["urls"]:
            short = u[:40] + ("…" if len(u) > 40 else "")
            contact_html += tag(f"🔗 {short}", "tag tag-green")
        if not contact_html:
            contact_html = tag("⚠ No contact info detected", "tag tag-red")

        st.markdown(contact_html, unsafe_allow_html=True)

        # Power verbs
        verbs = score_data["power_verbs"]
        if verbs:
            st.markdown("<br>", unsafe_allow_html=True)
            st.markdown('<div class="section-heading">Action Verbs Used</div>',
                        unsafe_allow_html=True)
            st.markdown(
                " ".join(tag(v.capitalize()) for v in verbs),
                unsafe_allow_html=True)

    # ── Tab 2: Skills ──
    with tab2:
        if not skills:
            st.warning("No skills detected. Make sure your resume lists specific tools, "
                       "languages, and technologies.")
        else:
            st.markdown('<div class="section-heading">Skills by Category</div>',
                        unsafe_allow_html=True)
            st.plotly_chart(chart_skill_distribution(skills),
                            use_container_width=True)

            for cat, skill_list in skills.items():
                with st.expander(f"**{cat}** — {len(skill_list)} skill(s)"):
                    st.markdown(
                        " ".join(f'<span class="tag">{s}</span>'
                                 for s in skill_list),
                        unsafe_allow_html=True)

    # ── Tab 3: Sections ──
    with tab3:
        st.markdown('<div class="section-heading">Resume Sections</div>',
                    unsafe_allow_html=True)

        section_icons = {
            "education":    "🎓",
            "experience":   "💼",
            "skills":       "🛠",
            "projects":     "🚀",
            "achievements": "🏆",
            "summary":      "📝",
            "contact":      "📬",
        }

        cols = st.columns(3)
        for i, (sec, present) in enumerate(sections.items()):
            icon = section_icons.get(sec, "📄")
            bg   = "#162520" if present else "#1c1020"
            bd   = "#2dd4bf" if present else "#f87171"
            lbl  = "Present" if present else "Missing"
            lbl_col = "#2dd4bf" if present else "#f87171"

            with cols[i % 3]:
                st.markdown(f"""
                <div class="card" style="border-left:4px solid {bd};
                                         background:{bg};text-align:center;
                                         padding:1.2rem 0.8rem;">
                  <div style="font-size:2rem">{icon}</div>
                  <div style="font-family:'Syne',sans-serif;font-weight:700;
                               font-size:0.9rem;margin:0.4rem 0 0.2rem">
                    {sec.title()}
                  </div>
                  <span class="tag" style="border-color:{lbl_col};
                                           color:{lbl_col};
                                           background:transparent;">
                    {lbl}
                  </span>
                </div>
                """, unsafe_allow_html=True)

    # ── Tab 4: Suggestions ──
    with tab4:
        high    = [s for s in suggestions if s["priority"] == "high"]
        medium  = [s for s in suggestions if s["priority"] == "medium"]
        good    = [s for s in suggestions if s["priority"] == "good"]

        if high:
            st.markdown('<div class="section-heading">🔴 High Priority</div>',
                        unsafe_allow_html=True)
            st.markdown('<div class="card card-accent">', unsafe_allow_html=True)
            for sug in high:
                st.markdown(
                    f'<div class="suggestion-item">'
                    f'<span class="sug-icon">{sug["icon"]}</span>'
                    f'<span>{sug["text"]}</span></div>',
                    unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        if medium:
            st.markdown('<div class="section-heading">🟡 Medium Priority</div>',
                        unsafe_allow_html=True)
            st.markdown('<div class="card">', unsafe_allow_html=True)
            for sug in medium:
                st.markdown(
                    f'<div class="suggestion-item">'
                    f'<span class="sug-icon">{sug["icon"]}</span>'
                    f'<span>{sug["text"]}</span></div>',
                    unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        if good:
            st.markdown('<div class="section-heading">✅ Strengths</div>',
                        unsafe_allow_html=True)
            st.markdown('<div class="card">', unsafe_allow_html=True)
            for sug in good:
                st.markdown(
                    f'<div class="suggestion-item">'
                    f'<span class="sug-icon">{sug["icon"]}</span>'
                    f'<span>{sug["text"]}</span></div>',
                    unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        # Download report button
        st.markdown("<br>", unsafe_allow_html=True)
        with st.spinner("Building PDF report…"):
            report_bytes = generate_pdf_report(
                uploaded.name, score_data, sections, skills, suggestions
            )
        st.markdown(
            get_download_link(
                report_bytes,
                f"resume_report_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                "⬇  Download Full Report (PDF)"
            ),
            unsafe_allow_html=True,
        )

    # ── Tab 5: Raw Text ──
    with tab5:
        st.markdown('<div class="section-heading">Extracted Text</div>',
                    unsafe_allow_html=True)
        st.markdown(
            f'<div class="scrollbox">{text[:4000]}'
            f'{"…[truncated]" if len(text) > 4000 else ""}</div>',
            unsafe_allow_html=True)
        st.caption(f"Total characters: {len(text):,} · Words: {wc:,} · "
                   f"Estimated pages: {pg}")

# ──────────────────────────────────────────────────────────────
#  PAGE: HOW IT WORKS
# ──────────────────────────────────────────────────────────────

def page_how_it_works():
    st.markdown("""
    <div class="hero">
      <p class="hero-title">How It Works</p>
      <p class="hero-sub">
        A transparent look at the analysis pipeline.
      </p>
    </div>
    """, unsafe_allow_html=True)

    steps = [
        ("1", "📤", "Upload", "Drop a PDF or DOCX resume into the uploader."),
        ("2", "🔍", "Parse",
         "pdfplumber (for PDF) or python-docx (for DOCX) extracts clean text."),
        ("3", "🏷", "Section Detection",
         "Keyword matching identifies headings like Education, Experience, Skills, Projects."),
        ("4", "🛠", "Skill Extraction",
         "~100 skills across 7 categories are matched with regex word-boundary search."),
        ("5", "🧮", "Scoring",
         "Five dimensions (sections, skills, contact, action verbs, detail) are scored "
         "out of 100 with transparent weights."),
        ("6", "💡", "Suggestions",
         "Rule-based engine generates prioritised tips — missing sections, weak verbs, "
         "lack of quantification, etc."),
        ("7", "📊", "Visualisation",
         "Plotly gauge, radar and bar charts render the analysis interactively."),
        ("8", "⬇", "Download",
         "ReportLab builds a polished PDF report you can save or share."),
    ]

    for num, icon, title, desc in steps:
        st.markdown(f"""
        <div class="card card-accent" style="display:flex;gap:1rem;align-items:flex-start;">
          <div style="font-family:'Syne',sans-serif;font-size:1.5rem;
                      font-weight:800;color:#2dd4bf;min-width:2rem">{num}</div>
          <div>
            <div style="font-family:'Syne',sans-serif;font-weight:700;
                        font-size:1rem;">{icon} {title}</div>
            <div style="font-size:0.88rem;color:#8b949e;margin-top:4px">{desc}</div>
          </div>
        </div>
        """, unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────
#  PAGE: ABOUT
# ──────────────────────────────────────────────────────────────

def page_about():
    st.markdown("""
    <div class="hero">
      <p class="hero-title">About This Project</p>
      <p class="hero-sub">
        Built for the FOSSEE Summer Internship Screening Task.
      </p>
    </div>

    <div class="card card-accent">
      <div class="section-heading">Project</div>
      <p style="font-size:0.9rem;color:#8b949e;">
        <b style="color:#e6edf3">AI Resume Analyzer</b> — an S6 Computer Science Engineering
        project demonstrating Python skills and UI/UX design using Streamlit.
        The app parses PDF/DOCX resumes, extracts skills, detects sections, scores
        the resume, and provides actionable feedback — all without any heavyweight ML models.
      </p>
    </div>

    <div class="card">
      <div class="section-heading">Tech Stack</div>
      <div>
        <span class="tag">Python 3.10+</span>
        <span class="tag">Streamlit</span>
        <span class="tag">pdfplumber</span>
        <span class="tag">python-docx</span>
        <span class="tag">Plotly</span>
        <span class="tag">ReportLab</span>
      </div>
    </div>

    <div class="card">
      <div class="section-heading">Design Principles</div>
      <p style="font-size:0.88rem;color:#8b949e;line-height:1.8">
        • Dark editorial theme with a teal accent palette<br>
        • Syne (display) + DM Sans (body) typography<br>
        • Structured tabs instead of a scrolling wall of content<br>
        • Colour-coded priority for suggestions (red / amber / green)<br>
        • Interactive Plotly charts — gauge, radar, horizontal bar<br>
        • Graceful error handling at every parse step
      </p>
    </div>
    """, unsafe_allow_html=True)

# ──────────────────────────────────────────────────────────────
#  ENTRY POINT
# ──────────────────────────────────────────────────────────────

def main():
    page = render_sidebar()

    if "Analyzer" in page:
        page_analyzer()
    elif "How It Works" in page:
        page_how_it_works()
    elif "About" in page:
        page_about()


if __name__ == "__main__":
    main()
